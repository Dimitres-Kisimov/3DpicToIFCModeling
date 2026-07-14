from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(3)

# --- Helper functions ---
def add_heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def add_paragraph(text, bold=False, italic=False, size=11, spacing_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    return p

def add_table_row(table, col1, col2, header=False):
    row = table.add_row()
    for i, text in enumerate([col1, col2]):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            run.font.bold = header
            run.font.size = Pt(10)
    return row

def shade_row(row, hex_color="D9E1F2"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("From 2D Image to IFC Model:\nA Critical Analysis of AI-Based 3D Reconstruction\nfor BIM Applications")
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Research Findings, Pipeline Evaluation, and Recommended Approaches")
run2.font.size = Pt(13)
run2.font.italic = True
run2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Dimitres Kisimov\n{datetime.date.today().strftime('%B %d, %Y')}\n3DpicToIFCModeling Project")
meta.runs[0].font.size = Pt(11)

doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading("Abstract", level=1)
add_paragraph(
    "This paper critically examines the 3DpicToIFCModeling project, which attempted to convert 2D photographs "
    "into IFC (Industry Foundation Classes) files suitable for Building Information Modeling (BIM) applications "
    "using AI-based single-view 3D reconstruction models. The project integrated multiple state-of-the-art models "
    "including InstantMesh, StableFast3D, TripoSR, TRELLIS, and Hunyuan3D-2, alongside mesh processing pipelines "
    "and a custom IFC export layer. Through systematic testing and evaluation, we identify fundamental misalignments "
    "between the capabilities of these AI models and the requirements of BIM-grade geometry. We document the findings, "
    "analyze the root causes of failure, and propose alternative approaches that are better suited to the stated goal."
)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading("Table of Contents", level=1)
toc_items = [
    "1. Introduction",
    "2. Project Overview and Pipeline",
    "3. Models Evaluated",
    "4. Critical Findings — What Went Wrong",
    "   4.1 Wrong AI Models for the Task",
    "   4.2 Single-View Reconstruction is Fundamentally Limited",
    "   4.3 Mesh Output is Incompatible with IFC Requirements",
    "   4.4 Absence of Background Segmentation",
    "   4.5 No Semantic Understanding",
    "5. Visual Evidence",
    "6. Why This Approach Cannot Work",
    "7. Recommended Alternative Approaches",
    "8. Lessons Learned",
    "9. Conclusion",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(11)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
add_heading("1. Introduction", level=1)
add_paragraph(
    "Building Information Modeling (BIM) represents a paradigm shift in the architecture, engineering, and construction "
    "(AEC) industry. IFC — the open standard format for BIM data — encodes not just geometry but semantic metadata: "
    "wall types, structural elements, material properties, spatial relationships, and object classifications. The appeal "
    "of automatically generating BIM models from photographs is enormous, promising to dramatically reduce the time "
    "and expertise required to digitize existing buildings or objects."
)
add_paragraph(
    "The 3DpicToIFCModeling project was born from this vision: take a single 2D photograph, run it through an AI model, "
    "and produce a valid IFC file. The project was developed across 8 phases, integrating multiple AI models, a Node.js "
    "backend, Python processing scripts, and a xeokit-based 3D viewer frontend. However, testing revealed severe "
    "quality issues with the outputs, prompting a fundamental reassessment of the approach."
)

# ============================================================
# 2. PROJECT OVERVIEW
# ============================================================
add_heading("2. Project Overview and Pipeline", level=1)
add_paragraph(
    "The pipeline was designed as a sequential, modular system consisting of five stages:"
)
stages = [
    ("Stage 1 — Image Input", "User uploads a 2D photograph via the web frontend."),
    ("Stage 2 — AI Model Selection", "User selects one of five available AI models: InstantMesh, StableFast3D, TripoSR, TRELLIS, or Hunyuan3D-2."),
    ("Stage 3 — 3D Mesh Generation", "The selected Python-based AI model processes the image and outputs a 3D mesh (typically OBJ or GLB format)."),
    ("Stage 4 — Mesh Processing", "Post-processing scripts clean, normalize, fix orientation, and convert the mesh to GLB."),
    ("Stage 5 — IFC Export", "A custom Python script wraps the mesh geometry in IFC4 format and triggers a download."),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0]
hdr.cells[0].text = "Stage"
hdr.cells[1].text = "Description"
for cell in hdr.cells:
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)
shade_row(hdr, "1F497D")
for cell in hdr.cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for stage, desc in stages:
    row = table.add_row()
    row.cells[0].text = stage
    row.cells[1].text = desc
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(
    "The system was functional from an engineering standpoint — the server ran, models were invoked, meshes were generated, "
    "and IFC files were produced. The failure was not in the implementation but in the fundamental suitability of the approach."
)

# ============================================================
# 3. MODELS EVALUATED
# ============================================================
add_heading("3. Models Evaluated", level=1)
models = [
    ("InstantMesh", "Tsinghua University", "Fast single-view mesh reconstruction using multi-view diffusion and instant-NGP. Optimized for speed over quality."),
    ("StableFast3D", "Stability AI", "Stable Diffusion-based 3D generation. Produces textured meshes but with significant hallucination artifacts."),
    ("TripoSR", "Tripo AI / Stability AI", "Transformer-based large reconstruction model. Higher quality but still single-view with inherent limitations."),
    ("TRELLIS", "MIT", "SLAT (Structured Latent) diffusion model. Produces structured outputs but still trained on general objects."),
    ("Hunyuan3D-2", "Tencent", "Multi-view aware texture generation on top of 3D shapes. Best texture quality but same geometric limitations."),
]
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0]
for i, h in enumerate(["Model", "Origin", "Description"]):
    hdr2.cells[i].text = h
    for run in hdr2.cells[i].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)
shade_row(hdr2, "1F497D")
for cell in hdr2.cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for name, origin, desc in models:
    row = table2.add_row()
    row.cells[0].text = name
    row.cells[1].text = origin
    row.cells[2].text = desc
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)

doc.add_paragraph()

# ============================================================
# 4. CRITICAL FINDINGS
# ============================================================
add_heading("4. Critical Findings — What Went Wrong", level=1)
add_paragraph(
    "Testing across all models and pipeline configurations revealed consistent, systematic failures. These are not bugs "
    "to be fixed — they are fundamental incompatibilities between what the technology can do and what BIM requires."
)

add_heading("4.1 Wrong AI Models for the Task", level=2)
add_paragraph(
    "All five models integrated into the pipeline are trained on ShapeNet, Objaverse, or similar datasets of everyday "
    "consumer objects: chairs, cars, animals, household items. Their internal representations and loss functions optimize "
    "for visual plausibility of general objects. They have never been trained on architectural geometry, structural elements, "
    "or BIM-relevant objects. Feeding them photographs and expecting BIM-grade output is analogous to using an image "
    "recognition model trained on cats to identify surgical instruments — the domain mismatch is categorical."
)

add_heading("4.2 Single-View Reconstruction is Fundamentally Limited", level=2)
add_paragraph(
    "The core technical problem is that 3D reconstruction from a single 2D image is a severely ill-posed problem. "
    "A single photograph provides no information about:"
)
bullets_42 = [
    "The geometry of occluded surfaces (back, bottom, sides not visible in the image)",
    "Absolute scale or real-world dimensions",
    "Depth relationships between objects in the scene",
    "Surface topology of non-visible areas",
]
for b in bullets_42:
    add_bullet(b)
add_paragraph(
    "Models like TripoSR handle this by hallucinating the missing geometry based on training data priors. For a toy car or "
    "a chair photographed in isolation, this produces visually acceptable results. For a building facade, a room, or a "
    "structural assembly, the hallucinated geometry is meaningless and metrically incorrect."
)

add_heading("4.3 Mesh Output is Incompatible with IFC Requirements", level=2)
add_paragraph(
    "IFC is not a mesh format. It is a semantic data model. An IFC file for a wall does not store a triangle mesh — "
    "it stores the wall's parametric definition: length, height, thickness, material layer set, boundary conditions, "
    "and spatial containment within a building storey. The approach of wrapping a noisy AI-generated triangle mesh "
    "in IFC XML produces a file that is technically parseable but semantically empty and geometrically useless for:"
)
bullets_43 = [
    "Structural analysis (no material properties, no load-bearing semantics)",
    "Energy simulation (no thermal zones, no envelope definition)",
    "Quantity take-off (no parametric dimensions)",
    "Clash detection (incorrect geometry)",
    "Code compliance checking (no semantic object types)",
]
for b in bullets_43:
    add_bullet(b)

add_heading("4.4 Absence of Background Segmentation", level=2)
add_paragraph(
    "A direct visual observation from testing: the 3D models generated by the pipeline included the image background "
    "as part of the reconstructed geometry. This manifests as a flat plane or curved surface surrounding the target "
    "object in the 3D viewer. This occurs because the AI models receive the full image without foreground/background "
    "separation. The TripoSR-SAM2-Humphrey-Enhanced branch attempted to address this with SAM2 segmentation, "
    "which improved isolation of the foreground object, but did not resolve the deeper geometric quality issues."
)

add_heading("4.5 No Semantic Understanding", level=2)
add_paragraph(
    "Even if the geometric quality were perfect, the pipeline has no mechanism to answer the questions IFC requires: "
    "Is this a wall or a column? What is its fire rating? Which storey does it belong to? What are its load-bearing "
    "properties? The AI models produce geometry only — no labels, no classifications, no relationships. "
    "Assigning IFC entity types (IfcWall, IfcSlab, IfcBeam) requires either manual annotation or a separate "
    "semantic segmentation and classification pipeline that was never built."
)

# ============================================================
# 5. VISUAL EVIDENCE
# ============================================================
add_heading("5. Visual Evidence", level=1)
add_paragraph(
    "During testing on the phase-2-sprints branch, a photograph was processed through the InstantMesh pipeline. "
    "The result displayed in the xeokit 3D viewer exhibited the following observable defects:"
)
evidence = [
    "Dark, monochromatic rendering with no material differentiation",
    "Inclusion of the background plane as part of the 3D geometry",
    "Floating mesh artifacts disconnected from the main body",
    "Jagged, high-frequency noise along mesh edges",
    "No recognizable correspondence to the original photograph subject",
    "Non-manifold geometry unsuitable for any downstream processing",
]
for e in evidence:
    add_bullet(e)
add_paragraph(
    "This output is representative of all models tested. The variation between models was in the degree of failure, "
    "not in whether failure occurred. TripoSR produced slightly cleaner topology; TRELLIS produced more structured "
    "outputs; but none produced geometry of sufficient quality or semantic richness for BIM use."
)

# ============================================================
# 6. WHY THIS APPROACH CANNOT WORK
# ============================================================
add_heading("6. Why This Approach Cannot Work", level=1)
add_paragraph(
    "The fundamental incompatibility can be summarized as a mismatch across three dimensions:"
)
table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0]
for i, h in enumerate(["Dimension", "What the AI Provides", "What BIM Requires"]):
    hdr3.cells[i].text = h
    for run in hdr3.cells[i].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(10)
shade_row(hdr3, "C00000")
for cell in hdr3.cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows3 = [
    ("Geometry", "Noisy triangle mesh, hallucinated surfaces, no scale", "Parametric, dimensionally accurate, manifold geometry"),
    ("Semantics", "None — raw geometry only", "Object types, properties, relationships, classifications"),
    ("Accuracy", "Visually plausible, metrically meaningless", "Survey-grade accuracy, real-world dimensions"),
    ("Data Model", "OBJ/GLB mesh file", "IFC entity graph with typed objects and spatial hierarchy"),
    ("Domain", "General consumer objects", "Architectural, structural, MEP elements"),
]
for r in rows3:
    row = table3.add_row()
    for i, text in enumerate(r):
        row.cells[i].text = text
        for run in row.cells[i].paragraphs[0].runs:
            run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(
    "The gap between these columns is not bridgeable by better AI models alone. It requires a fundamentally different "
    "pipeline architecture that starts from different inputs and uses different processing strategies."
)

# ============================================================
# 7. RECOMMENDED ALTERNATIVES
# ============================================================
add_heading("7. Recommended Alternative Approaches", level=1)

add_heading("7.1 Photogrammetry for Existing Structures", level=2)
add_paragraph(
    "For digitizing existing buildings or objects, photogrammetry using multiple photographs (20–200+ images) combined "
    "with Structure from Motion (SfM) algorithms produces metrically accurate point clouds. Tools such as Agisoft Metashape, "
    "RealityCapture, or open-source alternatives like COLMAP can process these into dense point clouds or meshes. "
    "These can then be registered into BIM software (Revit, ArchiCAD) as reference geometry for manual or semi-automated "
    "IFC model creation."
)

add_heading("7.2 LiDAR / Structured Light Scanning", level=2)
add_paragraph(
    "For survey-grade BIM, laser scanning (LiDAR) provides millimeter-accurate point clouds. The resulting scan-to-BIM "
    "workflow — now supported natively in Autodesk Revit, Trimble Connect, and others — produces IFC models with "
    "real-world dimensions. Mobile LiDAR (iPhone Pro, iPad Pro, Leica BLK series) has dramatically reduced the cost "
    "of entry for this workflow."
)

add_heading("7.3 Parametric Template Matching", level=2)
add_paragraph(
    "For objects with known typologies (standard doors, windows, furniture families), a more tractable approach is "
    "AI-based classification of the object type from a photograph, followed by retrieval of a parametric IFC template "
    "from a library. The user adjusts dimensions to match the photograph. This produces valid BIM data without requiring "
    "geometric reconstruction."
)

add_heading("7.4 Semantic Segmentation + Parametric Reconstruction", level=2)
add_paragraph(
    "For architectural floor plans or facade photographs, recent research demonstrates that semantic segmentation "
    "(identifying walls, openings, floors) followed by vectorization and parametric reconstruction can produce "
    "BIM-ready geometry. Projects such as RoomFormer, HouseGAN++, and various wall detection pipelines represent "
    "the current state of the art in this direction."
)

add_heading("7.5 Hybrid AI-Assisted BIM Authoring", level=2)
add_paragraph(
    "The most practical near-term approach is to use AI as an assistant within a BIM authoring tool rather than as "
    "a replacement for it. The AI helps with object recognition, dimension estimation, and model population, while "
    "a BIM professional validates and finalizes the IFC output. This hybrid approach is already being commercialized "
    "by companies such as Reconstruct, Alice Technologies, and Autodesk's AI-assisted design tools."
)

# ============================================================
# 8. LESSONS LEARNED
# ============================================================
add_heading("8. Lessons Learned", level=1)
lessons = [
    ("Define output requirements first", "The IFC standard requirements should have been the starting constraint, not an afterthought. Starting from the output format and working backwards would have immediately revealed the incompatibility."),
    ("Domain specificity matters in AI", "General-purpose 3D reconstruction models cannot be repurposed for domain-specific applications without domain-specific training data and loss functions."),
    ("Single-view reconstruction has a ceiling", "No matter how sophisticated, single-view reconstruction cannot recover information that is not present in the input. Multi-view or depth sensing is required for metrically useful outputs."),
    ("Semantic gap is the hardest problem", "Bridging raw geometry to semantic BIM objects requires explicit classification and relationship inference — a separate, non-trivial pipeline component."),
    ("Validate with domain experts early", "AEC domain experts (architects, structural engineers, BIM managers) should have been consulted during the design phase to validate whether the approach could meet professional standards."),
]
for title, body in lessons:
    p = doc.add_paragraph()
    run_bold = p.add_run(f"{title}: ")
    run_bold.font.bold = True
    run_bold.font.size = Pt(11)
    run_normal = p.add_run(body)
    run_normal.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ============================================================
# 9. CONCLUSION
# ============================================================
add_heading("9. Conclusion", level=1)
add_paragraph(
    "The 3DpicToIFCModeling project successfully demonstrated the technical feasibility of building an end-to-end "
    "pipeline connecting a web frontend to AI-based 3D reconstruction models and an IFC export layer. The engineering "
    "work — Node.js server, Python bridge, xeokit viewer, multi-model support — was competently executed."
)
add_paragraph(
    "However, the core premise — that single-view AI reconstruction models can produce BIM-grade IFC geometry from "
    "photographs — is fundamentally flawed. The models are trained on the wrong data, produce the wrong type of output, "
    "and cannot generate the semantic information that IFC requires. The outputs observed during testing confirm this "
    "analysis: noisy, artifact-laden meshes with no correspondence to real-world geometry, wrapped in IFC XML that "
    "provides no BIM value."
)
add_paragraph(
    "The path forward requires either a change of inputs (multi-view photogrammetry, LiDAR scanning) or a change of "
    "approach (parametric template matching, semantic segmentation pipelines, hybrid AI-assisted BIM authoring). "
    "The engineering infrastructure built during this project — particularly the web frontend, the Python processing "
    "bridge, and the IFC export layer — can be repurposed as components of a corrected architecture."
)
add_paragraph(
    "This project serves as a valuable case study in the importance of aligning AI model capabilities with application "
    "domain requirements before committing to a pipeline architecture. The lesson is generalizable: "
    "AI models must be evaluated not only for technical performance on benchmark datasets but for fitness-for-purpose "
    "within the specific constraints and standards of the target application domain.",
    italic=True
)

# ============================================================
# SAVE
# ============================================================
output_path = r"c:\Users\dimik\3DpicToIFCModeling\Research_Paper_3D_to_IFC_Critical_Analysis.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
