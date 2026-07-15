"""make_item_register.py — the ITEM LOGIC REGISTER: all 38 catalog categories,
each with its distribution logic, purpose, and standing vs the German standard.

Single source of truth below, three renderings:
  docs/ITEM_LOGIC_REGISTER.md                       (repo — GitHub renders tables)
  deliverable/local_only/Item_Logic_Register.docx   (local Word, real tables)
  frontend/item_register.html                       (in-app page, research hub)

    python scripts/make_item_register.py
"""
from pathlib import Path

REPO = Path(__file__).resolve().parents[1]

INTRO = (
    "How distribution works: the room type + area suggest the item list (counts grow "
    "with space; Dense staffs offices toward the ASR legal floor). Placement then runs "
    "in five layers: special pre-passes claim their spots first (presentation front "
    "zone, door flanks), the CP-SAT solver places all free-standing items with "
    "per-item clearances and preferences (wall / corner / open-centre), companion "
    "items attach to their parents (in-front / on-top / beside / ring), and wall "
    "decor mounts at human heights. Anything that cannot sit legally is refused and "
    "reported — never forced."
)

ASR_NOTE = (
    "German-standard status (confirmed): no rule below conflicts with ASR. Values "
    "marked ASR are cited from the legal text (verbatim quotes live in rule_packs.py, "
    "verified against the BGN-hosted full texts); values marked STRICTER exceed the "
    "legal minimum on purpose; values marked PRACTICE are ergonomic design rules "
    "(Neufert, Panero, modern workplace standards) that ASR does not regulate — they "
    "operate inside the ASR envelope, never against it."
)

FOOTER = (
    "Cross-cutting rules: clearances are per-item (default 0.15 m; seating 0.10; "
    "storage 0.20; the movable waste bin 0.05) and unknown/custom categories get a "
    "footprint-derived gap automatically. Routes follow ASR A1.8 by occupancy "
    "(0.90 / 1.00 / 1.20 m) with a 0.90 m person-path walked to every placed item — "
    "items only reachable through narrower gaps are reported UNREACHABLE. Door swings "
    "keep 0.90-1.2 m clear. Density tiers (Light / Medium / Dense) scale counts but "
    "never breach the legal envelope. Every companion placement is exported as a "
    "machine-checked link (in_front_of, on_top_of, beside, ring_around, throws_onto, "
    "audience_row_facing, door_flank, mounted_on, faces)."
)

STD = {  # item prefix -> standards basis
    "desk": "ASR A1.2 §5(3) staffing + 5.1.1/5.1.2 Bewegungsflaeche (cited)",
    "office chair": "inside the ASR workstation zone",
    "monitor": "PRACTICE (screen ergonomics)",
    "laptop": "PRACTICE (screen ergonomics)",
    "lamp (desk": "PRACTICE",
    "waste bin": "PRACTICE; clearance honesty",
    "partition": "PRACTICE (Grossraum acoustics)",
    "table (": "PRACTICE; ASR routes still enforced around it",
    "chair (": "PRACTICE (Neufert seating circulation)",
    "stool": "PRACTICE",
    "sofa": "PRACTICE",
    "coffee table": "PRACTICE (Panero 0.46 m)",
    "armchair": "PRACTICE",
    "side table": "PRACTICE",
    "cabinet": "PRACTICE; keeps ASR routes clear",
    "bookshelf": "PRACTICE",
    "filing cabinet": "PRACTICE",
    "locker": "PRACTICE",
    "server rack": "PRACTICE",
    "fridge": "PRACTICE",
    "microwave": "PRACTICE (working height)",
    "coffee machine": "PRACTICE",
    "water dispenser": "workplace welfare (ArbStaettV spirit)",
    "printer": "PRACTICE",
    "presentation screen": "PRACTICE (DIN-style ergonomics)",
    "whiteboard": "PRACTICE",
    "projector": "ASR A1.8 headroom >=2.00 m — ours 2.20 (STRICTER)",
    "lectern": "PRACTICE",
    "flipchart": "PRACTICE",
    "chair rows": "row pitch 0.90 m = ASR A1.8 Tab.2 (0.875 rounded STRICTER); viewing distance PRACTICE",
    "coat rack": "PRACTICE",
    "fire extinguisher": "ASR A2.2; grip 1.00 m inside the 0.80-1.20 m band",
    "first-aid cabinet": "presence per ASR A4.3; height PRACTICE (eye level)",
    "phone booth": "PRACTICE",
    "planter": "PRACTICE; keeps ASR routes clear",
    "lamp (floor": "PRACTICE",
    "bed": "residential — ASR n/a (A1.8 envelope applied anyway)",
    "mirror": "PRACTICE",
    "clock": "PRACTICE",
    "picture frame": "PRACTICE",
}


def _std(item):
    low = item.lower()
    for k, v in STD.items():
        if low.startswith(k):
            return v
    return "PRACTICE"


# (section title, [(item, distribution logic, purpose/meaning)])
SECTIONS = [
    ("1 · Workstation cluster — where work happens", [
        ("Desk (1.40×0.70 m)",
         "Backs toward walls; reserves the ASR Bewegungsflaeche in front (>=1.5 m², "
         ">=1.0 m deep — scales with desk width)",
         "The workstation itself; anchor of the whole cluster; count = ASR staffing rules"),
        ("Office chair",
         "In front of its desk, tucked into the reserved zone, rotated to face it",
         "The worker's seat; proves the desk is usable, not decorative"),
        ("Monitor",
         "On the desk surface, screen rotated toward the chair",
         "Screen work; the facing rule = a human must be able to read it"),
        ("Laptop",
         "On the desk, hinge at the rear, screen to the sitter",
         "Same as monitor; second surface slot"),
        ("Lamp (desk context)",
         "Third surface slot on the desk",
         "Task lighting at the workplace"),
        ("Waste bin",
         "Beside the desk, edge 0.10 m off the side panel; only a 0.05 m halo because "
         "it is EASILY MOVABLE (user rule); no desk in room -> opposite door flank "
         "from the coat rack",
         "Arm's-reach disposal; the flank rule keeps coats away from garbage"),
        ("Partition (divider)",
         "One per ~4 desks in open plans, screening between desk pairs",
         "Visual/acoustic privacy in Grossraum offices"),
    ]),
    ("2 · Social rings — where people gather", [
        ("Table (dining/work)",
         "Prefers OPEN floor when ringed (social pieces do not hide at walls)",
         "The gathering point; hosts rings and appliances"),
        ("Chair (regular, 0.45×0.52 m)",
         "Rings its table radially at chord-safe spacing, facing it, skipping any "
         "occupied front sector; never used as an office chair (user rule)",
         "Dining/meeting seating — human bodies around a shared surface"),
        ("Stool",
         "Rings the coffee table or table, same petal maths",
         "Casual perching; fills the ring where chairs would crowd"),
        ("Sofa",
         "Back flush to a wall (perimeter set)",
         "The lounge core; couches stay FIRST in every living list (user-approved)"),
        ("Coffee table",
         "0.46 m in front of the sofa (the '18-inch rule' — shin clearance)",
         "Within reach of a seated person"),
        ("Armchair",
         "Perimeter-leaning; gets a side-table companion",
         "Solo seating — reading corners, breakout zones, reception waiting"),
        ("Side table",
         "Directly beside its armchair",
         "A cup/book within arm's reach of whoever sits"),
    ]),
    ("3 · Storage — the walls' occupants", [
        ("Cabinet (1.2×0.6×1.8 m)",
         "Back flush to wall; count scales ~1 per 22 m²",
         "General storage / wardrobe; wall placement keeps floor routes free"),
        ("Bookshelf", "Wall-flush, ~1 per 30 m²", "Reference storage"),
        ("Filing cabinet", "Wall-flush, workspace-heavy rooms",
         "Document storage — the 'workspace' room type's identity"),
        ("Locker", "Wall-flush, ~1 per 26-34 m²",
         "Personal storage (German offices expect them)"),
        ("Server rack", "Wall-flush, workspace >52 m²",
         "IT infrastructure for technical workrooms"),
        ("Fridge", "Wall-flush perimeter, kitchens + break rooms >14 m²",
         "Cold storage; belongs to kitchen / Pausenraum only (user rule)"),
    ]),
    ("4 · Appliances — on surfaces, never the floor", [
        ("Microwave",
         "ON the table/counter (on_top link, elev ~0.74 m) in kitchens and break "
         "rooms — never floor-placed (user rule)",
         "Heating food at working height"),
        ("Coffee machine", "On the table in break rooms / kitchens",
         "The social magnet of every Pausenraum"),
        ("Water dispenser", "Perimeter; offices >50 m², break rooms >12 m²",
         "Hydration point (workplace welfare)"),
        ("Printer",
         "On its stand (never the ground — user rule), wall-side, 1 per ~8 "
         "workstations, second >75 m²",
         "Shared MFP; central so everyone reaches it"),
    ]),
    ("5 · Presentation kit — the lecture geometry", [
        ("Presentation screen", "Front wall, centre, mounted at 0.80 m",
         "The projection surface everyone must see"),
        ("Whiteboard",
         "Front wall left of the screen, 0.90 m; becomes the projector's target "
         "only when no screen exists",
         "Writing surface at standing-arm height"),
        ("Projector",
         "CEILING at 2.2 m (ASR headroom kept below), throw distance 1.2× image "
         "width, lens aimed at the display — in ANY room type",
         "Projects onto its display; the aim is a machine-checked throws_onto link"),
        ("Lectern", "Front zone, rotated 180° to face the audience",
         "The speaker's station"),
        ("Flipchart", "Beside the lectern", "The speaker's second surface, at hand"),
        ("Chair rows (audience)",
         "Parallel rows CENTRED on the display axis (partial last row too), facing "
         "it; first row at ~1.5× image width; row pitch >=0.90 m (ASR aisle)",
         "Everyone sees the screen at legal viewing/egress spacing (user rule: "
         "'never forget this')"),
    ]),
    ("6 · Entry & safety — the door's neighbourhood and the law's items", [
        ("Coat rack", "Wall immediately left or right of the door",
         "Where a human hangs a coat on entry (user rule)"),
        ("Fire extinguisher", "Wall-mounted 1.00 m; offices >40 m²",
         "ASR A2.2 spirit — grabbing height, visible"),
        ("First-aid cabinet", "Wall-mounted 1.35 m; rooms >50 m² / break rooms",
         "Eye-level emergency access"),
        ("Phone booth", "Free-standing box; offices >62 m²",
         "Call privacy in open plans (modern workplace standard)"),
    ]),
    ("7 · Ambience & residential", [
        ("Planter",
         "CORNERS (solver drives both nearest-wall gaps to zero) or queued along "
         "the wall when corners fill — never mid-room (user rule)",
         "Greenery without stealing circulation"),
        ("Lamp (floor)", "Free accent near seating",
         "Ambient light in living/quiet rooms"),
        ("Bed (1.60×2.05 m)",
         "Listed FIRST in bedrooms (essentials-first: the solver keeps it over "
         "accents); needs one long side reachable",
         "The bedroom's reason to exist"),
        ("Mirror", "Left wall, face height 1.50 m",
         "Dressing check — full length at human eye line"),
        ("Clock",
         "2.05 m high on the wall the SEATED person faces (computed from the "
         "chair's gaze)",
         "Readable from the working position"),
        ("Picture frame", "Wall at 1.55 m eye level, avoiding doors and other decor",
         "Decor at gallery height"),
    ]),
]


def write_markdown():
    out = ["# Item Logic Register — every item, its distribution logic, its purpose",
           "", INTRO, "", "> " + ASR_NOTE, ""]
    for title, rows in SECTIONS:
        out += ["## " + title, "",
                "| Item | Distribution logic | Purpose / meaning | Standard basis |",
                "|---|---|---|---|"]
        out += ["| **%s** | %s | %s | %s |" % (i, l, p, _std(i)) for i, l, p in rows]
        out += [""]
    out += ["## Cross-cutting rules", "", FOOTER, ""]
    p = REPO / "docs" / "ITEM_LOGIC_REGISTER.md"
    p.write_text("\n".join(out), encoding="utf-8")
    print("wrote", p)


def write_docx():
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading("Item Logic Register", level=0)
    doc.add_paragraph("Every catalog item, its distribution logic, its purpose, and "
                      "its standing vs the German standard — exactly as the SCS "
                      "Studio engine executes it.")
    doc.add_paragraph(INTRO)
    doc.add_paragraph(ASR_NOTE)
    for title, rows in SECTIONS:
        doc.add_heading(title, level=1)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text = "Item", "Distribution logic"
        hdr[2].text, hdr[3].text = "Purpose / meaning", "Standard basis"
        for i, l, p in rows:
            c = t.add_row().cells
            c[0].text, c[1].text, c[2].text, c[3].text = i, l, p, _std(i)
        for row in t.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
    doc.add_heading("Cross-cutting rules", level=1)
    doc.add_paragraph(FOOTER)
    p = REPO / "deliverable" / "local_only" / "Item_Logic_Register.docx"
    p.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(p))
    except PermissionError:            # the file is open in Word — save alongside
        p = p.with_name("Item_Logic_Register_updated.docx")
        doc.save(str(p))
    print("wrote", p)


def write_html():
    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;")
    parts = []
    for title, rows in SECTIONS:
        parts.append("<h2>%s</h2>" % esc(title))
        parts.append("<div class='tw'><table><tr><th>Item</th><th>Distribution logic"
                     "</th><th>Purpose / meaning</th><th>Standard basis</th></tr>")
        for i, l, p in rows:
            std = _std(i)
            cls = ("asr" if ("ASR" in std or "ArbStaett" in std) and "STRICTER" not in std
                   else ("str" if "STRICTER" in std else "pr"))
            parts.append("<tr><td><b>%s</b></td><td>%s</td><td>%s</td>"
                         "<td class='%s'>%s</td></tr>" % (esc(i), esc(l), esc(p), cls, esc(std)))
        parts.append("</table></div>")
    html = """<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Item Logic Register — distribution rules, in tables</title>
<link rel="icon" href="data:image/svg+xml,<svg xmlns=%%22http://www.w3.org/2000/svg%%22 viewBox=%%220 0 100 100%%22><text y=%%22.9em%%22 font-size=%%2290%%22>📋</text></svg>">
<style>
* { margin:0; padding:0; box-sizing:border-box; }
:root { --bg:#f4f6fb; --panel:#fff; --ink:#1f2733; --mut:#6b7688; --line:#e3e8f0; --brand:#2f6bff; }
body { font-family:'Segoe UI',system-ui,sans-serif; background:var(--bg); color:var(--ink); }
header { background:var(--panel); border-bottom:1px solid var(--line); padding:16px 28px;
         display:flex; align-items:baseline; gap:14px; position:sticky; top:0; z-index:5; }
header h1 { font-size:19px; } header p { color:var(--mut); font-size:13px; }
header a { margin-left:auto; color:var(--brand); font-size:13px; text-decoration:none; font-weight:600; }
main { max-width:1240px; margin:0 auto; padding:22px 28px 70px; }
.note { background:var(--panel); border:1px solid var(--line); border-left:4px solid var(--brand);
        border-radius:12px; padding:14px 18px; font-size:13.5px; line-height:1.65; margin-bottom:8px; }
h2 { font-size:15px; margin:26px 2px 8px; }
.tw { overflow-x:auto; background:var(--panel); border:1px solid var(--line); border-radius:12px; }
table { border-collapse:collapse; width:100%%; font-size:12.5px; }
th, td { text-align:left; padding:8px 12px; border-bottom:1px solid var(--line); vertical-align:top; }
th { background:#f0f4fc; font-size:11.5px; text-transform:uppercase; letter-spacing:.04em; color:var(--mut); }
td.asr { color:#1a7f4b; font-weight:600; } td.str { color:#1f52d6; font-weight:600; }
td.pr { color:var(--mut); }
.legend b { font-weight:700; }
</style></head><body>
<header><h1>📋 Item Logic Register</h1>
<p>every item, its distribution logic, its purpose — and its standing vs the German standard</p>
<a href="/hub.html">← Research hub</a></header>
<main>
<div class="note"><b>%s</b></div>
<div class="note legend">%s<br>
<span style="color:#1a7f4b;font-weight:700">green = ASR-cited</span> ·
<span style="color:#1f52d6;font-weight:700">blue = stricter than ASR</span> ·
<span style="color:#6b7688">grey = ergonomic practice (never against ASR)</span></div>
%s
<div class="note">%s</div>
</main></body></html>""" % (esc(INTRO), esc(ASR_NOTE), "\n".join(parts), esc(FOOTER))
    p = REPO / "frontend" / "item_register.html"
    p.write_text(html, encoding="utf-8")
    print("wrote", p)


if __name__ == "__main__":
    write_markdown()
    write_docx()
    write_html()
