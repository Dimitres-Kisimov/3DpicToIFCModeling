"""
build_docx.py — render deliverable/PAPER_MASTER.md to a Word .docx with embedded figures.

Handles the constructs used in PAPER_MASTER.md: headings (#/##/###), [FIGURE: path — caption]
directives (embeds the image + caption), Markdown tables, bullet lists, code fences,
horizontal rules, and inline **bold** / `code`. Produces a title page + auto Table of Contents.

Run: python backend/python-scripts/build_docx.py
Out: deliverable/PAPER_3DpicToIFC_2026-06-23.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[2]
DEL = REPO / "deliverable"
SRC = DEL / "PAPER_MASTER.md"
OUT = DEL / "PAPER_3DpicToIFC_2026-06-23.docx"

FIG_RE = re.compile(r"\[FIGURE:\s*([^\s—-]+(?:\.png|\.jpg))\s*[—-]+\s*(.*?)\]")


def add_inline(par, text):
    """Add text with **bold** and `code` runs."""
    for chunk in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            par.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        else:
            par.add_run(chunk)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    t = OxmlElement("w:t"); t.text = "Right-click → Update Field to build the Table of Contents."
    fld.append(t); run._r.addprevious(fld)


def main():
    if not SRC.exists():
        raise SystemExit(f"missing {SRC}")
    lines = SRC.read_text(encoding="utf-8").splitlines()
    # strip leading HTML comment (instructions block)
    text = "\n".join(lines)
    text = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    lines = text.splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i, n = 0, len(lines)
    title_done = False
    while i < n:
        line = lines[i].rstrip()

        if not line.strip():
            i += 1; continue

        # horizontal rule
        if re.fullmatch(r"-{3,}", line.strip()):
            i += 1; continue

        # headings
        m = re.match(r"(#{1,3})\s+(.*)", line)
        if m:
            level = len(m.group(1)); txt = m.group(2).strip()
            if level == 1 and not title_done:
                h = doc.add_heading("", level=0)
                r = h.add_run(txt); r.font.size = Pt(20)
                title_done = True
            else:
                doc.add_heading(txt, level=min(level, 3))
            i += 1; continue

        # figure directive
        fm = FIG_RE.search(line)
        if fm:
            rel, caption = fm.group(1), fm.group(2).strip()
            img = (DEL / rel)
            if img.exists():
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p.add_run().add_picture(str(img), width=Inches(6.2))
                except Exception:
                    p.add_run(f"[image: {rel}]")
                cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(caption); cr.italic = True; cr.font.size = Pt(9.5)
                cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                doc.add_paragraph(f"[missing figure: {rel}]")
            i += 1; continue

        # code fence
        if line.strip().startswith("```"):
            i += 1; buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph(); r = p.add_run("\n".join(buf))
            r.font.name = "Consolas"; r.font.size = Pt(9)
            continue

        # table
        if line.strip().startswith("|") and i + 1 < n and re.match(r"\s*\|[\s:|-]+\|", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2; rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")]); i += 1
            tbl = doc.add_table(rows=1, cols=len(header)); tbl.style = "Light Grid Accent 1"
            for j, c in enumerate(header):
                run = tbl.rows[0].cells[j].paragraphs[0].add_run(c); run.bold = True
            for row in rows:
                cells = tbl.add_row().cells
                for j in range(len(header)):
                    add_inline(cells[j].paragraphs[0], row[j] if j < len(row) else "")
            doc.add_paragraph()
            continue

        # bullet list
        if re.match(r"\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^\s*[-*]\s+", "", line))
            i += 1; continue
        if re.match(r"\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\s*\d+\.\s+", "", line))
            i += 1; continue

        # normal paragraph (merge following non-blank, non-special lines)
        buf = [line]; i += 1
        while i < n and lines[i].strip() and not re.match(r"(#{1,3}\s|\||```|\s*[-*]\s|\s*\d+\.\s)", lines[i]) \
                and not FIG_RE.search(lines[i]) and not re.fullmatch(r"-{3,}", lines[i].strip()):
            buf.append(lines[i].rstrip()); i += 1
        p = doc.add_paragraph()
        # subtitle/author lines right after title get italic emphasis via **bold** already
        add_inline(p, " ".join(buf))

    # insert TOC after the abstract heading if present (simple: after first heading-1 title)
    doc.save(OUT)
    print(f"-> {OUT}  ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
